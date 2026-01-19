import streamlit as st
import requests
import json
import PyPDF2
import pandas as pd
import io
import re
from docx import Document

# Read API key from secrets.toml
api_key = st.secrets["AZURE_OPENAI_API_KEY"]

# Azure OpenAI configuration
AZURE_ENDPOINT = "https://mcqgpt.openai.azure.com"
AZURE_DEPLOYMENT_NAME = "o4-mini"
AZURE_API_VERSION = "2025-01-01-preview"

# Function to call Azure OpenAI API
def call_azure_openai_api(prompt):
    url = f"{AZURE_ENDPOINT}/openai/deployments/{AZURE_DEPLOYMENT_NAME}/chat/completions?api-version={AZURE_API_VERSION}"
    headers = {
        "Content-Type": "application/json",
        "api-key": api_key
    }
    body = {
        "messages": [
            {"role": "user", "content": prompt}
        ],
        "temperature": 0.8,
        "max_completion_tokens": 8192
    }
    
    try:
        response = requests.post(url, headers=headers, json=body)
        
        if response.status_code != 200:
            st.error(f"API Request failed with status code: {response.status_code}")
            st.error(f"Response: {response.text}")
            return None
        
        content = response.json()
        
        # Extract the generated text
        if "choices" in content and len(content["choices"]) > 0:
            return content["choices"][0]["message"]["content"]
        elif "error" in content:
            raise Exception(f"API Error: {content['error']}")
        else:
            raise Exception(f"Unexpected response format: {content}")
            
    except requests.exceptions.RequestException as e:
        raise Exception(f"Request failed: {str(e)}")
    except Exception as e:
        raise Exception(f"Error processing API response: {str(e)}")

# Function to extract text from PDF with page tracking
def extract_text_from_pdf(pdf_file):
    pdf_reader = PyPDF2.PdfReader(pdf_file)
    text = ""
    page_texts = []
    
    for page_num, page in enumerate(pdf_reader.pages):
        page_text = page.extract_text()
        text += f"[PAGE {page_num + 1}]\n{page_text}\n\n"
        page_texts.append((page_num + 1, page_text))
    
    return text, pdf_reader, page_texts

# Function to parse MCQs from AI response
def parse_mcqs(response_text):
    """
    Parse MCQs from the AI response into structured format.
    """
    mcqs = []
    
    # Split by numbered questions
    question_blocks = re.split(r'(?=\d+\.\s+)', response_text)
    
    for block in question_blocks:
        if not block.strip():
            continue
        
        # Extract question number and full text
        question_match = re.match(r'(\d+)\.\s+(.*?)(?=\n[A-E]\)|\Z)', block, re.DOTALL)
        if not question_match:
            continue
        
        question_num = question_match.group(1)
        question_text = question_match.group(2).strip()
        
        # Extract options A-E
        options = {}
        for letter in ['A', 'B', 'C', 'D', 'E']:
            pattern = f'{letter}\\)\s+(.*?)(?=\n[A-E]\\)|\nCorrect Answer|\nAnswer:|\n\n|\Z)'
            opt_match = re.search(pattern, block, re.DOTALL)
            if opt_match:
                options[letter] = opt_match.group(1).strip()
        
        # Extract correct answer
        answer_match = re.search(r'(?:Correct Answer|Answer):\s*([A-E])', block, re.IGNORECASE)
        correct_answer = answer_match.group(1).upper() if answer_match else None
        
        # Extract page reference if available
        page_match = re.search(r'(?:Page|From page)\s+(\d+)', block, re.IGNORECASE)
        page_num = page_match.group(1) if page_match else "N/A"
        
        if question_text and len(options) == 5 and correct_answer:
            mcqs.append({
                'number': question_num,
                'question': question_text,
                'options': options,
                'answer': correct_answer,
                'page': page_num
            })
    
    return mcqs

# Streamlit app
st.title("🦷 Restorative Dentistry MCQ Generator")
st.markdown("*Clinical scenario-based MCQ generation for prosthodontic education*")

# Sidebar configuration
with st.sidebar:
    st.header("⚙️ Configuration")
    
    cognition_level = st.selectbox(
        "Bloom's Cognitive Level",
        [
            "C1 - Knowledge (Recall)",
            "C2 - Comprehension (Understand)",
            "C3 - Application (Clinical relevance)",
            "C4 - Analysis (Critical thinking)",
            "C5 - Synthesis (Integrate)",
            "C6 - Evaluation (Judge)"
        ],
        index=2  # Default to C3
    )
    
    difficulty_level = st.selectbox(
        "Difficulty Level",
        ["Easy", "Moderate", "Hard"],
        index=1  # Default to Moderate
    )
    
    num_questions = st.slider(
        "Number of Questions",
        min_value=1,
        max_value=10,
        value=5
    )
    
    st.markdown("---")
    st.markdown("### About Cognitive Levels")
    st.info(
        "**C1-C2**: Basic recall and understanding\n\n"
        "**C3-C6**: Clinical application and critical thinking\n\n"
        "Recommended: C3+ for resident evaluation"
    )

# Main content
uploaded_file = st.file_uploader("📄 Upload PDF File", type=["pdf"])

custom_prompt = st.text_area(
    "Additional Instructions (Optional)",
    placeholder="E.g., Focus on fixed prosthodontics, include implant cases, emphasize material selection...",
    height=100
)

# Generate MCQ button
if st.button("🎯 Generate Clinical MCQs", type="primary"):
    if uploaded_file is not None:
        try:
            with st.spinner("📖 Extracting content from PDF..."):
                pdf_text, pdf_reader, page_texts = extract_text_from_pdf(uploaded_file)
                st.success(f"✅ PDF loaded: {len(pdf_reader.pages)} pages extracted")
            
            # Extract cognition level code
            cog_level = cognition_level.split(" ")[0]
            
            # Construct the enhanced prompt with your specific requirements
            prompt = f'''You are an experienced PROSTHODONTIST creating MCQs to evaluate prosthodontic residents. Generate exactly {num_questions} Multiple Choice Questions based on the following PDF content.

PDF CONTENT:
{pdf_text}

CRITICAL REQUIREMENTS:

1. COGNITIVE FRAMEWORK (Bloom's Taxonomy):
   - Knowledge: Basic recall of facts
   - Comprehension: Understanding concepts
   - Application: Using knowledge in clinical situations
   - Analysis: Breaking down complex problems
   - Synthesis: Integrating multiple concepts
   - Evaluation: Making clinical judgments
   
   Target Level: {cognition_level}
   Each MCQ must assess clinical relevance and critical thinking appropriate to this level.

2. MCQ STRUCTURE:
   - Start with a complete clinical scenario that mirrors real prosthodontic practice
   - Follow with a leading question that naturally flows from the scenario
   - Provide exactly 5 options (A through E)
   - All options must be similarly formatted, relevant, and plausible
   - Only ONE option should be the best or correct answer
   - Avoid obvious distractors or joke options

3. WRITING STYLE:
   - Write like a human prosthodontist, not a textbook
   - Use clear, direct, professional but conversational language
   - No buzzwords, no press release tone, no em dashes
   - Keep it natural, as if you're discussing a case with a colleague
   - Use dental educational terminology appropriately

4. DIFFICULTY: {difficulty_level}

{f"5. ADDITIONAL FOCUS: {custom_prompt}" if custom_prompt else ""}

FORMAT EACH QUESTION EXACTLY AS FOLLOWS:

1. [Complete clinical scenario describing patient presentation, history, findings, etc.]

What is the most appropriate next step/diagnosis/treatment/material choice?

A) [Option A - realistic and plausible]
B) [Option B - realistic and plausible]
C) [Option C - realistic and plausible]
D) [Option D - realistic and plausible]
E) [Option E - realistic and plausible]

Correct Answer: [A/B/C/D/E]

---

[Repeat this format for all {num_questions} questions]

Remember: These MCQs should evaluate a resident's ability to think critically and apply prosthodontic principles in realistic clinical situations.'''
            
            # Call Azure OpenAI API
            with st.spinner("🤖 Generating clinical scenario MCQs..."):
                response = call_azure_openai_api(prompt)
            
            if response:
                st.success("✅ MCQs generated successfully!")
                
                # Display raw response in expander for review
                with st.expander("📋 View Full AI Response"):
                    st.text(response)
                
                # Parse MCQs
                parsed_mcqs = parse_mcqs(response)
                
                if len(parsed_mcqs) == 0:
                    st.error("❌ Failed to parse MCQs. Please check the AI response format.")
                    st.write("Raw response:")
                    st.write(response)
                elif len(parsed_mcqs) < num_questions:
                    st.warning(f"⚠️ Only {len(parsed_mcqs)} out of {num_questions} questions were successfully parsed.")
                
                if parsed_mcqs:
                    # Display MCQs in a nice format
                    st.markdown("---")
                    st.subheader("📚 Generated MCQs")
                    
                    for mcq in parsed_mcqs:
                        with st.container():
                            st.markdown(f"**Question {mcq['number']}:**")
                            st.write(mcq['question'])
                            st.markdown("")
                            for letter in ['A', 'B', 'C', 'D', 'E']:
                                st.write(f"{letter}) {mcq['options'][letter]}")
                            st.markdown(f"**Correct Answer:** :green[{mcq['answer']}]")
                            if mcq['page'] != "N/A":
                                st.caption(f"📄 Reference: Page {mcq['page']}")
                            st.markdown("---")
                    
                    # Create DataFrame for export
                    mcq_data = pd.DataFrame({
                        "Question Number": [mcq['number'] for mcq in parsed_mcqs],
                        "Clinical Scenario & Question": [mcq['question'] for mcq in parsed_mcqs],
                        "Option A": [mcq['options']['A'] for mcq in parsed_mcqs],
                        "Option B": [mcq['options']['B'] for mcq in parsed_mcqs],
                        "Option C": [mcq['options']['C'] for mcq in parsed_mcqs],
                        "Option D": [mcq['options']['D'] for mcq in parsed_mcqs],
                        "Option E": [mcq['options']['E'] for mcq in parsed_mcqs],
                        "Correct Answer": [mcq['answer'] for mcq in parsed_mcqs],
                        "Cognitive Level": [cognition_level] * len(parsed_mcqs),
                        "Difficulty Level": [difficulty_level] * len(parsed_mcqs),
                        "Page Reference": [mcq['page'] for mcq in parsed_mcqs],
                        "Source Document": [uploaded_file.name] * len(parsed_mcqs)
                    })
                    
                    # Display table
                    st.subheader("📊 MCQ Summary Table")
                    st.dataframe(mcq_data, use_container_width=True)
                    
                    # Download options
                    col1, col2 = st.columns(2)
                    
                    with col1:
                        # CSV Download
                        csv = mcq_data.to_csv(index=False).encode("utf-8")
                        st.download_button(
                            label="📥 Download as CSV",
                            data=csv,
                            file_name=f"mcqs_{uploaded_file.name.replace('.pdf', '')}.csv",
                            mime="text/csv"
                        )
                    
                    with col2:
                        # Word Download
                        doc = Document()
                        doc.add_heading("Restorative Dentistry MCQs", level=1)
                        doc.add_paragraph(f"Source: {uploaded_file.name}")
                        doc.add_paragraph(f"Cognitive Level: {cognition_level}")
                        doc.add_paragraph(f"Difficulty: {difficulty_level}")
                        doc.add_paragraph("")
                        
                        for mcq in parsed_mcqs:
                            doc.add_heading(f"Question {mcq['number']}", level=2)
                            doc.add_paragraph(mcq['question'])
                            doc.add_paragraph("")
                            
                            for letter in ['A', 'B', 'C', 'D', 'E']:
                                doc.add_paragraph(f"{letter}) {mcq['options'][letter]}")
                            
                            doc.add_paragraph("")
                            doc.add_paragraph(f"Correct Answer: {mcq['answer']}")
                            if mcq['page'] != "N/A":
                                doc.add_paragraph(f"Page Reference: {mcq['page']}")
                            doc.add_paragraph("")
                            doc.add_paragraph("_" * 80)
                            doc.add_paragraph("")
                        
                        doc_bytes = io.BytesIO()
                        doc.save(doc_bytes)
                        doc_bytes.seek(0)
                        
                        st.download_button(
                            label="📥 Download as Word",
                            data=doc_bytes,
                            file_name=f"mcqs_{uploaded_file.name.replace('.pdf', '')}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
            
        except Exception as e:
            st.error(f"❌ Error: {str(e)}")
            st.exception(e)
    else:
        st.warning("⚠️ Please upload a PDF file to begin.")

# New MCQ button
if st.button("🔄 Generate New Set"):
    st.rerun()

# Footer
st.markdown("---")
col1, col2 = st.columns(2)
with col1:
    st.markdown(
        '<div style="text-align: center;">Powered by <a href="https://medentec.com/" target="_blank">Medentec</a></div>',
        unsafe_allow_html=True
    )
with col2:
    st.markdown(
        '<div style="text-align: center;">Code by <a href="https://www.linkedin.com/public-profile/settings?trk=d_flagship3_profile_self_view_public_profile" target="_blank">Dr. Fahad Umer</a></div>',
        unsafe_allow_html=True
    )
